"""
RAG Backend v4.0 — 全面升级版
改进：
  - 异步文档处理（后台任务）
  - 支持 DOCX / XLSX / CSV 文档格式
  - 会话持久化（JSON 存储）
  - 流式 SSE 更稳定（心跳 + 超时保护）
  - 向量检索结果附带相关度分数
  - 知识库统计信息增强
  - 启动时自动创建默认知识库
  - 更健壮的错误处理 + 详细日志
  - /health 接口返回服务状态详情
  - 文档上传支持进度跟踪（后台任务）
启动: cd backend && uvicorn main:app --reload --port 8000
"""

import os
import uuid
import json
import base64
import asyncio
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
from datetime import datetime
from contextlib import asynccontextmanager
from concurrent.futures import ThreadPoolExecutor

import httpx
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from dotenv import load_dotenv
load_dotenv(Path(__file__).parent / ".env")

from langchain_openai import ChatOpenAI
from langchain_community.document_loaders import PyPDFLoader, TextLoader, WebBaseLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.embeddings import Embeddings
from openai import OpenAI as OpenAIClient

# ==================== 日志配置 ====================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("rag")

# ==================== 配置 ====================

API_KEY = os.getenv("DASHSCOPE_API_KEY")
if not API_KEY:
    raise ValueError("请设置环境变量 DASHSCOPE_API_KEY")

BASE_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1"
LLM_MODEL = os.getenv("LLM_MODEL", "qwen-plus")
EMB_MODEL = os.getenv("EMB_MODEL", "text-embedding-v3")
VL_MODEL = os.getenv("VL_MODEL", "qwen-vl-plus")

BASE_DIR = Path(__file__).parent
UPLOAD_DIR = BASE_DIR / "uploads"
INDEX_DIR = BASE_DIR / "faiss_index"
SESSION_DIR = BASE_DIR / "sessions"

for d in [UPLOAD_DIR, INDEX_DIR, SESSION_DIR]:
    d.mkdir(exist_ok=True)

_http_sync = httpx.Client(timeout=120.0)
_http_async = httpx.AsyncClient(timeout=120.0)
_executor = ThreadPoolExecutor(max_workers=4)

# 后台任务状态跟踪
upload_tasks: Dict[str, Dict] = {}

# ==================== 自定义 Embedding ====================

class DashScopeEmbeddings(Embeddings):
    def __init__(self, model: str = "text-embedding-v3", batch_size: int = 10):
        self.model = model
        self.batch_size = batch_size
        self.client = OpenAIClient(api_key=API_KEY, base_url=BASE_URL)

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        results = []
        for i in range(0, len(texts), self.batch_size):
            batch = texts[i:i + self.batch_size]
            results.extend(self._embed_batch(batch))
        return results

    def embed_query(self, text: str) -> List[float]:
        response = self.client.embeddings.create(model=self.model, input=text)
        return response.data[0].embedding

    def _embed_batch(self, texts: List[str]) -> List[List[float]]:
        response = self.client.embeddings.create(model=self.model, input=texts)
        return [item.embedding for item in response.data]


embeddings = DashScopeEmbeddings(model=EMB_MODEL)
llm = ChatOpenAI(
    model=LLM_MODEL,
    base_url=BASE_URL,
    api_key=API_KEY,
    temperature=0.2,
    streaming=True,
    http_client=_http_sync,
)

# ==================== 查询扩展与检索工具 ====================

def expand_query(question: str, history_context: str = "") -> List[str]:
    """查询扩展：生成多视角检索查询，提升召回率。"""
    try:
        client = OpenAIClient(api_key=API_KEY, base_url=BASE_URL)
        history_hint = f"\n对话历史摘要：{history_context}" if history_context else ""
        prompt = f"""你是检索优化专家。基于用户问题，生成2个语义相关但角度不同的检索查询。
用户问题：{question}{history_hint}
要求：每行一个，简洁（15字以内），覆盖不同角度，只输出查询本身。
输出2个查询："""
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=100,
        )
        expanded = [q.strip() for q in response.choices[0].message.content.strip().split("\n") if q.strip()][:2]
        return [question] + expanded
    except Exception as e:
        logger.warning(f"查询扩展失败: {e}")
        return [question]


def deduplicate_docs(docs: List[Document], threshold: float = 0.82) -> List[Document]:
    """基于 Jaccard 相似度去重。"""
    def jaccard(a: str, b: str) -> float:
        sa, sb = set(a.split()), set(b.split())
        return len(sa & sb) / len(sa | sb) if sa and sb else 0.0
    unique = []
    for doc in docs:
        if not any(jaccard(doc.page_content, u.page_content) > threshold for u in unique):
            unique.append(doc)
    return unique


def score_relevance(query: str, docs: List[Document]) -> List[Tuple[Document, float]]:
    """关键词匹配补充打分。"""
    keywords = set(query.lower().split())
    scored = []
    for doc in docs:
        content_lower = doc.page_content.lower()
        hit = sum(1 for kw in keywords if kw in content_lower)
        scored.append((doc, hit / max(len(keywords), 1)))
    scored.sort(key=lambda x: x[1], reverse=True)
    return scored


# ==================== 文档加载器扩展 ====================

def load_docx(path: str) -> List[Document]:
    """加载 DOCX 文档。"""
    try:
        import docx
        doc = docx.Document(path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [Document(page_content=text, metadata={"source": path, "type": "docx"})]
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")


def load_xlsx(path: str) -> List[Document]:
    """加载 XLSX / CSV 文档，每个 sheet 或行为一段。"""
    try:
        import pandas as pd
        ext = Path(path).suffix.lower()
        if ext == ".csv":
            df = pd.read_csv(path)
            text = df.to_markdown(index=False)
            return [Document(page_content=text, metadata={"source": path, "type": "csv"})]
        else:
            xl = pd.ExcelFile(path)
            docs = []
            for sheet in xl.sheet_names:
                df = xl.parse(sheet)
                text = f"[Sheet: {sheet}]\n" + df.to_markdown(index=False)
                docs.append(Document(page_content=text, metadata={"source": path, "sheet": sheet, "type": "xlsx"}))
            return docs
    except ImportError:
        raise ImportError("请安装 pandas + openpyxl: pip install pandas openpyxl tabulate")


def load_image_ocr(path: str, filename: str) -> List[Document]:
    """用视觉模型 OCR 提取图片文字。"""
    ext = Path(path).suffix.lower()
    with open(path, "rb") as f:
        img_base64 = base64.b64encode(f.read()).decode()
    prompt = "请识别并完整输出图片中的所有文字内容，保持原有结构，包括表格、公式和标注。"
    response = _http_sync.post(
        f"{BASE_URL}/chat/completions",
        headers={"Authorization": f"Bearer {API_KEY}"},
        json={
            "model": VL_MODEL,
            "messages": [{"role": "user", "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/{ext[1:]};base64,{img_base64}"}}
            ]}],
            "max_tokens": 2000
        }
    )
    result = response.json()
    text_content = result["choices"][0]["message"]["content"] if "choices" in result else str(result)
    return [Document(page_content=text_content, metadata={"source": filename, "type": "image"})]


# ==================== 知识库管理器 ====================

class KnowledgeBase:
    def __init__(self):
        self.vectorstores: Dict[str, FAISS] = {}
        self.metadata: Dict[str, Dict] = {}
        self._load_all()

    def _load_all(self):
        for kb_dir in INDEX_DIR.iterdir():
            if kb_dir.is_dir():
                kb_id = kb_dir.name
                try:
                    vs = FAISS.load_local(str(kb_dir), embeddings, allow_dangerous_deserialization=True)
                    self.vectorstores[kb_id] = vs
                    meta_file = kb_dir / "meta.json"
                    if meta_file.exists():
                        with open(meta_file) as f:
                            self.metadata[kb_id] = json.load(f)
                    logger.info(f"已加载知识库: {kb_id}")
                except Exception as e:
                    logger.error(f"加载知识库 {kb_id} 失败: {e}")

    def create(self, kb_id: str, name: str, description: str = "") -> bool:
        if kb_id in self.vectorstores:
            return False
        self.vectorstores[kb_id] = None
        self.metadata[kb_id] = {
            "name": name,
            "description": description,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "doc_count": 0,
            "chunk_count": 0,
            "sources": [],
            "file_types": {},
        }
        self._save_meta(kb_id)
        logger.info(f"已创建知识库: {kb_id} ({name})")
        return True

    def delete(self, kb_id: str) -> bool:
        if kb_id not in self.vectorstores:
            return False
        del self.vectorstores[kb_id]
        del self.metadata[kb_id]
        kb_dir = INDEX_DIR / kb_id
        if kb_dir.exists():
            import shutil
            shutil.rmtree(kb_dir)
        logger.info(f"已删除知识库: {kb_id}")
        return True

    def rename(self, kb_id: str, new_name: str, new_description: Optional[str] = None) -> bool:
        if kb_id not in self.metadata:
            return False
        self.metadata[kb_id]["name"] = new_name
        if new_description is not None:
            self.metadata[kb_id]["description"] = new_description
        self._save_meta(kb_id)
        return True

    def add_documents(self, kb_id: str, docs: List[Document], splitter_config: Dict = None):
        if kb_id not in self.vectorstores:
            raise ValueError(f"知识库 {kb_id} 不存在")

        cfg = splitter_config or {"chunk_size": 800, "chunk_overlap": 200}
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=cfg.get("chunk_size", 800),
            chunk_overlap=cfg.get("chunk_overlap", 200),
            length_function=len,
            separators=["\n\n\n", "\n\n", "\n", "。", "！", "？", ". ", "，", ""],
        )
        chunks = splitter.split_documents(docs)

        if not chunks:
            logger.warning(f"文档分块结果为空，kb_id={kb_id}")
            return 0

        if self.vectorstores[kb_id] is None:
            self.vectorstores[kb_id] = FAISS.from_documents(chunks, embeddings)
        else:
            self.vectorstores[kb_id].add_documents(chunks)

        self._save(kb_id)
        self._update_meta(kb_id, chunks, docs)
        logger.info(f"知识库 {kb_id} 添加了 {len(chunks)} 个片段（来自 {len(docs)} 个文档）")
        return len(chunks)

    def search_with_expansion(
        self,
        kb_id: str,
        query: str,
        history_context: str = "",
        k: int = 6,
    ) -> Tuple[List[Document], List[str], List[float]]:
        """多路召回 + 去重 + 相关性重排，返回 (docs, queries, scores)。"""
        if kb_id not in self.vectorstores or self.vectorstores[kb_id] is None:
            return [], [query], []
        vs = self.vectorstores[kb_id]
        queries = expand_query(query, history_context)

        all_docs: List[Document] = []
        for q in queries:
            try:
                docs = vs.max_marginal_relevance_search(q, k=k, fetch_k=k * 4, lambda_mult=0.6)
                all_docs.extend(docs)
            except Exception:
                docs = vs.similarity_search(q, k=k)
                all_docs.extend(docs)

        unique_docs = deduplicate_docs(all_docs, threshold=0.80)
        scored = score_relevance(query, unique_docs)
        final_docs = [doc for doc, _ in scored[:k]]
        final_scores = [round(score, 3) for _, score in scored[:k]]
        return final_docs, queries, final_scores

    def search(self, kb_id: str, query: str, k: int = 6) -> List[Document]:
        docs, _, _ = self.search_with_expansion(kb_id, query, k=k)
        return docs

    def _save(self, kb_id: str):
        kb_dir = INDEX_DIR / kb_id
        kb_dir.mkdir(exist_ok=True)
        self.vectorstores[kb_id].save_local(str(kb_dir))

    def _save_meta(self, kb_id: str):
        meta_file = INDEX_DIR / kb_id / "meta.json"
        meta_file.parent.mkdir(exist_ok=True)
        with open(meta_file, "w", encoding="utf-8") as f:
            json.dump(self.metadata[kb_id], f, ensure_ascii=False, indent=2)

    def _update_meta(self, kb_id: str, chunks: List[Document], raw_docs: List[Document]):
        meta = self.metadata[kb_id]
        sources = set(meta.get("sources", []))
        file_types: Dict[str, int] = meta.get("file_types", {})
        for doc in raw_docs:
            src = doc.metadata.get("source", "unknown")
            sources.add(src.split("/")[-1] if "/" in src else src)
            ftype = doc.metadata.get("type", Path(src).suffix.lstrip(".") or "unknown")
            file_types[ftype] = file_types.get(ftype, 0) + 1
        meta["sources"] = list(sources)
        meta["file_types"] = file_types
        meta["chunk_count"] = len(self.vectorstores[kb_id].docstore._dict)
        meta["doc_count"] = len(sources)
        meta["updated_at"] = datetime.now().isoformat()
        self._save_meta(kb_id)

    def get_info(self, kb_id: str) -> Optional[Dict]:
        if kb_id not in self.metadata:
            return None
        return {**self.metadata[kb_id], "id": kb_id, "has_index": self.vectorstores.get(kb_id) is not None}

    def list_all(self) -> List[Dict]:
        return [self.get_info(kb_id) for kb_id in self.vectorstores.keys()]

    def list_documents(self, kb_id: str) -> List[Dict]:
        if kb_id not in self.vectorstores or self.vectorstores[kb_id] is None:
            return []
        docs = []
        for doc_id, doc in self.vectorstores[kb_id].docstore._dict.items():
            docs.append({
                "id": doc_id,
                "content": doc.page_content[:200] + ("..." if len(doc.page_content) > 200 else ""),
                "source": doc.metadata.get("source", "未知"),
                "page": doc.metadata.get("page"),
                "type": doc.metadata.get("type", "text"),
                "char_count": len(doc.page_content),
            })
        return docs

    def delete_document(self, kb_id: str, doc_id: str) -> bool:
        if kb_id not in self.vectorstores or self.vectorstores[kb_id] is None:
            return False
        vs = self.vectorstores[kb_id]
        if doc_id not in vs.docstore._dict:
            return False
        all_ids = [i for i in vs.docstore._dict.keys() if i != doc_id]
        if all_ids:
            vs.delete([doc_id])
            self._update_meta(kb_id, [vs.docstore._dict[i] for i in all_ids], [vs.docstore._dict[i] for i in all_ids])
        else:
            self.vectorstores[kb_id] = None
            kb_dir = INDEX_DIR / kb_id
            if kb_dir.exists():
                for f in kb_dir.iterdir():
                    if f.name != "meta.json":
                        f.unlink()
            self.metadata[kb_id].update({"chunk_count": 0, "doc_count": 0, "sources": []})
            self._save_meta(kb_id)
        return True


# ==================== 会话管理（持久化）====================

class SessionManager:
    def __init__(self):
        self._sessions: Dict[str, Dict] = {}
        self._load_all()

    def _load_all(self):
        for f in SESSION_DIR.glob("*.json"):
            try:
                with open(f) as fp:
                    data = json.load(fp)
                sid = f.stem
                self._sessions[sid] = {
                    "kb_id": data.get("kb_id", "default"),
                    "history": [],  # 不持久化消息内容（隐私），仅恢复会话元数据
                    "created_at": data.get("created_at", datetime.now().isoformat()),
                }
            except Exception:
                pass

    def get_or_create(self, session_id: str, kb_id: str) -> Dict:
        if session_id not in self._sessions:
            self._sessions[session_id] = {
                "history": [],
                "kb_id": kb_id,
                "created_at": datetime.now().isoformat(),
            }
            self._save(session_id)
        else:
            self._sessions[session_id]["kb_id"] = kb_id
        return self._sessions[session_id]

    def clear(self, session_id: str):
        if session_id in self._sessions:
            self._sessions[session_id]["history"] = []
            self._save(session_id)

    def _save(self, session_id: str):
        data = {k: v for k, v in self._sessions[session_id].items() if k != "history"}
        with open(SESSION_DIR / f"{session_id}.json", "w") as f:
            json.dump(data, f)

    def __contains__(self, item):
        return item in self._sessions

    def __getitem__(self, item):
        return self._sessions[item]

    def __setitem__(self, key, value):
        self._sessions[key] = value


kb_manager = KnowledgeBase()
session_manager = SessionManager()

# ==================== FastAPI 应用 ====================

@asynccontextmanager
async def lifespan(app: FastAPI):
    # 启动时确保默认知识库存在
    if "default" not in kb_manager.vectorstores:
        kb_manager.create("default", "默认知识库", "系统自动创建的默认知识库")
        logger.info("已创建默认知识库")
    logger.info(f"RAG 系统 v4.0 启动完成 | LLM: {LLM_MODEL} | Embedding: {EMB_MODEL}")
    yield
    await _http_async.aclose()
    _http_sync.close()
    _executor.shutdown(wait=False)

app = FastAPI(
    title="RAG API v4",
    description="增强版 RAG 系统 — 查询扩展 + MMR 检索 + 相关性重排 + 多格式文档支持",
    version="4.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== Pydantic 模型 ====================

class CreateKBRequest(BaseModel):
    name: str
    description: str = ""

class RenameKBRequest(BaseModel):
    name: str
    description: Optional[str] = None

class ChatRequest(BaseModel):
    session_id: str
    question: str
    kb_id: str = "default"
    stream: bool = True

class AddUrlRequest(BaseModel):
    urls: List[str]
    kb_id: str = "default"
    splitter_config: Optional[Dict] = None

class SearchRequest(BaseModel):
    query: str
    kb_id: str = "default"
    top_k: int = 6

# ==================== 工具函数 ====================

def format_docs_with_sources(docs: List[Document], scores: List[float] = None) -> Tuple[str, List[Dict]]:
    parts, sources, seen = [], [], {}
    for i, d in enumerate(docs):
        src = d.metadata.get("source", "未知")
        page = d.metadata.get("page")
        src_label = src.split("/")[-1] if "/" in src else src
        if src not in seen:
            seen[src] = len(sources) + 1
            sources.append({
                "index": seen[src],
                "source": src,
                "label": src_label,
                "page": page,
                "score": scores[i] if scores else None,
                "type": d.metadata.get("type", "text"),
            })
        idx = seen[src]
        loc = f"[来源{idx}: {src_label}" + (f", 第{page+1}页" if page is not None else "") + "]"
        parts.append(f"{loc}\n{d.page_content}")
    return "\n\n---\n\n".join(parts), sources


SYSTEM_TEMPLATE = """你是专业的文档问答助手，基于检索到的文档片段精准回答问题。

## 回答规则
1. **优先使用文档内容**：充分整合所有检索片段，给出完整回答
2. **引用来源**：引用具体信息时用 [来源N] 标注（如 [来源1]）
3. **诚实回答**：文档中没有的信息请明确说明，不编造
4. **结构化表达**：复杂问题使用标题/列表组织回答
5. **语言**：统一使用中文

## 检索到的文档片段
{context}
---
请基于以上内容详细回答用户问题。"""


def build_history_context(history: List) -> str:
    if not history:
        return ""
    parts = []
    for msg in history[-4:]:
        role = "用户" if isinstance(msg, HumanMessage) else "助手"
        parts.append(f"{role}: {msg.content[:100]}")
    return " | ".join(parts)


# ==================== 后台文档处理 ====================

def _process_upload_task(task_id: str, kb_id: str, save_path: Path, filename: str, ext: str):
    """在线程池中处理文档，完成后更新任务状态。"""
    try:
        upload_tasks[task_id]["status"] = "processing"
        upload_tasks[task_id]["message"] = "正在解析文档…"

        if ext in (".png", ".jpg", ".jpeg", ".webp"):
            docs = load_image_ocr(str(save_path), filename)
        elif ext == ".pdf":
            loader = PyPDFLoader(str(save_path))
            docs = loader.load()
        elif ext == ".docx":
            docs = load_docx(str(save_path))
        elif ext in (".xlsx", ".xls"):
            docs = load_xlsx(str(save_path))
        elif ext == ".csv":
            docs = load_xlsx(str(save_path))
        elif ext == ".html":
            try:
                from langchain_community.document_loaders import UnstructuredHTMLLoader
                docs = UnstructuredHTMLLoader(str(save_path)).load()
            except Exception:
                docs = TextLoader(str(save_path), encoding="utf-8").load()
        else:
            docs = TextLoader(str(save_path), encoding="utf-8").load()

        upload_tasks[task_id]["message"] = f"正在向量化 {len(docs)} 段内容…"
        chunk_count = kb_manager.add_documents(kb_id, docs)

        upload_tasks[task_id].update({
            "status": "done",
            "message": f"成功添加 {chunk_count} 个片段",
            "chunk_count": chunk_count,
            "doc_count": len(docs),
        })
    except Exception as e:
        logger.error(f"处理文档失败 task={task_id}: {e}")
        upload_tasks[task_id].update({"status": "error", "message": str(e)})
    finally:
        # 清理临时文件
        try:
            save_path.unlink(missing_ok=True)
        except Exception:
            pass


# ==================== API 路由 ====================

@app.get("/")
def root():
    return {"status": "ok", "service": "RAG System", "version": "4.0.0"}

@app.get("/health")
def health():
    return {
        "status": "healthy",
        "version": "4.0.0",
        "llm_model": LLM_MODEL,
        "emb_model": EMB_MODEL,
        "kb_count": len(kb_manager.vectorstores),
        "session_count": len(session_manager._sessions),
        "pending_tasks": sum(1 for t in upload_tasks.values() if t["status"] == "processing"),
        "timestamp": datetime.now().isoformat(),
    }

# ----- 知识库管理 -----

@app.post("/kb/create")
def create_kb(req: CreateKBRequest):
    is_default = req.name in ("默认知识库", "default") and "default" not in kb_manager.vectorstores
    kb_id = "default" if is_default else uuid.uuid4().hex[:8]
    kb_manager.create(kb_id, req.name, req.description)
    return {"kb_id": kb_id, "message": "知识库创建成功"}

@app.get("/kb/list")
def list_kb():
    return {"knowledge_bases": kb_manager.list_all()}

@app.get("/kb/{kb_id}")
def get_kb(kb_id: str):
    info = kb_manager.get_info(kb_id)
    if not info:
        raise HTTPException(status_code=404, detail="知识库不存在")
    return info

@app.delete("/kb/{kb_id}")
def delete_kb(kb_id: str):
    if not kb_manager.delete(kb_id):
        raise HTTPException(status_code=404, detail="知识库不存在")
    return {"message": "知识库已删除"}

@app.put("/kb/{kb_id}")
def rename_kb(kb_id: str, req: RenameKBRequest):
    if not kb_manager.rename(kb_id, req.name, req.description):
        raise HTTPException(status_code=404, detail="知识库不存在")
    return {"message": "知识库已更新"}

# ----- 文档管理 -----

@app.get("/kb/{kb_id}/docs")
def list_docs(kb_id: str, page: int = Query(1, ge=1), page_size: int = Query(50, ge=1, le=200)):
    if kb_id not in kb_manager.vectorstores:
        raise HTTPException(status_code=404, detail="知识库不存在")
    all_docs = kb_manager.list_documents(kb_id)
    total = len(all_docs)
    start = (page - 1) * page_size
    end = start + page_size
    return {
        "documents": all_docs[start:end],
        "total": total,
        "page": page,
        "page_size": page_size,
        "pages": (total + page_size - 1) // page_size,
    }

@app.delete("/kb/{kb_id}/docs/{doc_id}")
def delete_doc(kb_id: str, doc_id: str):
    if kb_id not in kb_manager.vectorstores:
        raise HTTPException(status_code=404, detail="知识库不存在")
    if not kb_manager.delete_document(kb_id, doc_id):
        raise HTTPException(status_code=404, detail="文档不存在")
    return {"message": "文档已删除"}

# ----- 文件上传 -----

SUPPORTED_EXTS = (".pdf", ".txt", ".md", ".html", ".png", ".jpg", ".jpeg", ".webp", ".docx", ".xlsx", ".xls", ".csv")

@app.post("/upload")
async def upload_file(
    background_tasks: BackgroundTasks,
    kb_id: str = "default",
    file: UploadFile = File(...),
    async_mode: bool = Query(False, description="是否使用异步后台处理"),
):
    ext = Path(file.filename).suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的格式 {ext}，支持：{', '.join(SUPPORTED_EXTS)}"
        )

    if kb_id not in kb_manager.vectorstores:
        kb_manager.create(kb_id, kb_id, "自动创建")

    # 保存文件
    task_id = uuid.uuid4().hex[:8]
    save_path = UPLOAD_DIR / f"{task_id}{ext}"
    save_path.write_bytes(await file.read())

    if async_mode:
        upload_tasks[task_id] = {
            "status": "pending",
            "message": "等待处理…",
            "filename": file.filename,
            "kb_id": kb_id,
            "created_at": datetime.now().isoformat(),
        }
        background_tasks.add_task(
            _process_upload_task, task_id, kb_id, save_path, file.filename, ext
        )
        return {"task_id": task_id, "message": "文档已加入处理队列", "filename": file.filename}
    else:
        # 同步处理
        try:
            _process_upload_task(task_id, kb_id, save_path, file.filename, ext)
            task = upload_tasks.get(task_id, {})
            if task.get("status") == "error":
                raise HTTPException(status_code=500, detail=task.get("message", "处理失败"))
            return {
                "message": task.get("message", "上传成功"),
                "filename": file.filename,
                "kb_id": kb_id,
                "chunk_count": task.get("chunk_count", 0),
            }
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

@app.get("/upload/task/{task_id}")
def get_upload_task(task_id: str):
    if task_id not in upload_tasks:
        raise HTTPException(status_code=404, detail="任务不存在")
    return upload_tasks[task_id]

# ----- 添加网页 -----

@app.post("/add-url")
def add_url(req: AddUrlRequest):
    if req.kb_id not in kb_manager.vectorstores:
        kb_manager.create(req.kb_id, req.kb_id, "自动创建")
    all_docs, errors = [], []
    for url in req.urls:
        try:
            loader = WebBaseLoader(url)
            docs = loader.load()
            for doc in docs:
                doc.metadata["source"] = url
            all_docs.extend(docs)
        except Exception as e:
            errors.append(f"{url}: {str(e)}")
    chunk_count = 0
    if all_docs:
        chunk_count = kb_manager.add_documents(req.kb_id, all_docs, req.splitter_config)
    return {"message": f"已添加 {chunk_count} 个片段", "errors": errors, "kb_id": req.kb_id}

# ----- 搜索 -----

@app.post("/search")
def search(req: SearchRequest):
    docs, queries, scores = kb_manager.search_with_expansion(req.kb_id, req.query, k=req.top_k)
    _, sources = format_docs_with_sources(docs, scores)
    return {
        "results": [{"content": d.page_content, "metadata": d.metadata, "score": scores[i] if i < len(scores) else None} for i, d in enumerate(docs)],
        "sources": sources,
        "expanded_queries": queries,
    }

# ----- 对话（流式）-----

@app.post("/chat/stream")
def chat_stream(req: ChatRequest):
    if req.kb_id not in kb_manager.vectorstores or kb_manager.vectorstores[req.kb_id] is None:
        raise HTTPException(status_code=400, detail="知识库为空，请先上传文件")

    session = session_manager.get_or_create(req.session_id, req.kb_id)
    history = session["history"]

    async def generate():
        try:
            history_context = build_history_context(history)
            docs, expanded_queries, scores = kb_manager.search_with_expansion(
                req.kb_id, req.question, history_context, k=6
            )
            context, sources = format_docs_with_sources(docs, scores)

            yield f"data: {json.dumps({'type': 'sources', 'sources': sources, 'expanded_queries': expanded_queries}, ensure_ascii=False)}\n\n"

            messages = [{"role": "system", "content": SYSTEM_TEMPLATE.format(context=context)}]
            for msg in history[-6:]:
                if isinstance(msg, HumanMessage):
                    messages.append({"role": "user", "content": msg.content})
                elif isinstance(msg, AIMessage):
                    messages.append({"role": "assistant", "content": msg.content})
            messages.append({"role": "user", "content": req.question})

            client = OpenAIClient(api_key=API_KEY, base_url=BASE_URL)
            stream = client.chat.completions.create(
                model=LLM_MODEL,
                messages=messages,
                stream=True,
                temperature=0.2,
                max_tokens=2048,
            )

            full_response = ""
            # 心跳计数：每 20 个 token 发一次 keep-alive 防止连接超时
            heartbeat_counter = 0
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    content = chunk.choices[0].delta.content
                    full_response += content
                    yield f"data: {json.dumps({'type': 'content', 'content': content}, ensure_ascii=False)}\n\n"
                    heartbeat_counter += 1
                    if heartbeat_counter % 20 == 0:
                        yield ": keep-alive\n\n"

            # 更新历史
            history.append(HumanMessage(content=req.question))
            history.append(AIMessage(content=full_response))
            if len(history) > 20:
                session["history"] = history[-20:]

            yield f"data: {json.dumps({'type': 'done', 'total_chars': len(full_response)}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"流式对话错误: {e}")
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        }
    )

# ----- 对话（非流式）-----

@app.post("/chat")
def chat(req: ChatRequest):
    if req.kb_id not in kb_manager.vectorstores or kb_manager.vectorstores[req.kb_id] is None:
        raise HTTPException(status_code=400, detail="知识库为空")

    session = session_manager.get_or_create(req.session_id, req.kb_id)
    history = session["history"]

    try:
        history_context = build_history_context(history)
        docs, _, scores = kb_manager.search_with_expansion(req.kb_id, req.question, history_context, k=6)
        context, sources = format_docs_with_sources(docs, scores)

        messages = [{"role": "system", "content": SYSTEM_TEMPLATE.format(context=context)}]
        for msg in history[-6:]:
            if isinstance(msg, HumanMessage):
                messages.append({"role": "user", "content": msg.content})
            elif isinstance(msg, AIMessage):
                messages.append({"role": "assistant", "content": msg.content})
        messages.append({"role": "user", "content": req.question})

        client = OpenAIClient(api_key=API_KEY, base_url=BASE_URL)
        response = client.chat.completions.create(
            model=LLM_MODEL, messages=messages, temperature=0.2, max_tokens=2048
        )
        answer = response.choices[0].message.content

        history.append(HumanMessage(content=req.question))
        history.append(AIMessage(content=answer))
        if len(history) > 20:
            session["history"] = history[-20:]

        return {"answer": answer, "session_id": req.session_id, "sources": sources}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ----- 会话管理 -----

@app.delete("/session/{session_id}")
def clear_session(session_id: str):
    session_manager.clear(session_id)
    return {"message": "会话已清空"}

# ----- 兼容旧接口 -----

@app.delete("/knowledge/{kb_id}")
def clear_knowledge_legacy(kb_id: str):
    kb_manager.delete(kb_id)
    return {"message": "知识库已删除"}